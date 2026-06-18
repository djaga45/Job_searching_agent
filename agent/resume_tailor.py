from pathlib import Path

from docx import Document

from agent.config import TAILORED_DIR
from agent.llm import invoke_json
from agent.models import Experience, JobPosting, TailoredResume, UserProfile
from agent.profile_loader import find_master_resume, load_resume_text, profile_to_text


def tailor_resume(profile: UserProfile, job: JobPosting) -> tuple[TailoredResume, Path]:
    master_resume = find_master_resume()
    resume_context = (
        load_resume_text(master_resume) if master_resume else profile_to_text(profile)
    )

    system = (
        "You tailor resumes honestly. NEVER invent employers, degrees, or skills the "
        "candidate does not have. You may reorder bullets, rewrite wording, and "
        "emphasize relevant experience. Return strict JSON with keys: summary (string), "
        "experience (array of {title, company, dates, bullets}), skills_highlight "
        "(array of strings), changes_made (array of strings)."
    )
    user = (
        f"MASTER RESUME / PROFILE:\n{resume_context}\n\n"
        f"TARGET JOB:\nTitle: {job.title}\nCompany: {job.company}\n"
        f"Description:\n{job.description[:4000]}\n\n"
        "Tailor this resume for the job. Keep facts truthful."
    )

    result = invoke_json(system, user)
    tailored = TailoredResume(
        job_id=job.id,
        summary=str(result["summary"]),
        experience=[
            Experience(
                title=item["title"],
                company=item["company"],
                dates=item["dates"],
                bullets=list(item["bullets"]),
            )
            for item in result["experience"]
        ],
        skills_highlight=list(result.get("skills_highlight", [])),
        changes_made=list(result.get("changes_made", [])),
    )
    output_path = _write_docx(profile, tailored, job)
    return tailored, output_path


def _write_docx(profile: UserProfile, tailored: TailoredResume, job: JobPosting) -> Path:
    doc = Document()
    doc.add_heading(profile.name, level=0)
    contact_bits = [profile.email]
    if profile.phone:
        contact_bits.append(profile.phone)
    if profile.linkedin:
        contact_bits.append(profile.linkedin)
    if profile.location:
        contact_bits.append(profile.location)
    doc.add_paragraph(" | ".join(contact_bits))

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(tailored.summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(tailored.skills_highlight or profile.skills))

    doc.add_heading("Experience", level=1)
    for exp in tailored.experience:
        doc.add_paragraph(f"{exp.title} — {exp.company} ({exp.dates})")
        for bullet in exp.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    if profile.education:
        doc.add_heading("Education", level=1)
        for edu in profile.education:
            doc.add_paragraph(f"{edu.degree}, {edu.institution} ({edu.dates})")

    safe_company = "".join(c if c.isalnum() else "_" for c in job.company)[:30]
    safe_title = "".join(c if c.isalnum() else "_" for c in job.title)[:30]
    filename = f"{safe_company}_{safe_title}_{job.id}.docx"
    output_path = TAILORED_DIR / filename
    doc.save(output_path)
    return output_path